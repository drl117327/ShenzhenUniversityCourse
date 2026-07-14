"""
构建实验五最终报告 - 基于 WSL 真实数据
i9-14900HX, GCC 13.3.0, Ubuntu 24.04 WSL2
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = r'C:\Users\D1382\Desktop\experiment5_code'
OUTPUT = os.path.join(BASE_DIR, '实验五_实验报告_最终版.docx')

# 图片路径 (使用真实数据版本)
IMG = {
    'fig1_1': os.path.join(BASE_DIR, 'fig1_1_execution_time_real.png'),
    'fig1_2': os.path.join(BASE_DIR, 'fig1_2_speedup_real.png'),
    'fig2_1': os.path.join(BASE_DIR, 'fig2_1_mountain_3d_real.png'),
    'fig2_2': os.path.join(BASE_DIR, 'fig2_2_mountain_contour_real.png'),
    'fig2_3': os.path.join(BASE_DIR, 'fig2_3_throughput_vs_size_real.png'),
    'fig2_4': os.path.join(BASE_DIR, 'fig2_4_throughput_vs_stride_real.png'),
    'fig2_5': os.path.join(BASE_DIR, 'fig2_5_cache_verify_real.png'),
}

# ---- Helpers ----
def set_cell_shading(cell, color):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def add_code_block(doc, code, title=None):
    if title:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'))
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    doc.add_paragraph()

def add_img(doc, path, width=Inches(5.3), caption=None):
    if not os.path.exists(path):
        print(f'  WARN: missing {path}')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    doc.add_paragraph()

# ---- Build ----
def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'; style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ====== 封面 ======
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('深 圳 大 学 实 验 报 告').bold = True
    p.runs[0].font.size = Pt(22)
    doc.add_paragraph()

    table = doc.add_table(rows=8, cols=2); table.style = 'Table Grid'
    cover = [
        ('课程名称：','计算机系统(2)'),
        ('实验项目名称：','Cache实验'),
        ('学院：','计算机与软件学院'),
        ('专业：','人工智能卓越班'),
        ('指导教师：','罗秋明'),
        ('报告人：','邓瑞霖    学号：2024150040    班级：01'),
        ('实验时间：','2026年 6 月 4 日 -- 2026年 6 月 19 日'),
        ('实验报告提交时间：','2026年 6 月  日'),
    ]
    for i,(k,v) in enumerate(cover):
        c0=table.cell(i,0); c0.text=''; r0=c0.paragraphs[0].add_run(k)
        r0.bold=True; r0.font.size=Pt(12)
        c0.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        c1=table.cell(i,1); c1.text=''
        c1.paragraphs[0].add_run(v).font.size=Pt(12)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('教务处制').bold=True; p.runs[0].font.size=Pt(14)
    doc.add_page_break()

    # ====== 一、实验目的 ======
    doc.add_heading('一、实验目的',1)
    for t in ['加强对Cache工作原理的理解；',
              '体验程序中访存模式变化是如何影响Cache效率进而影响程序性能的过程；',
              '学习在X86真实机器上通过调整程序访存模式来探测多级Cache结构以及TLB的大小。']:
        doc.add_paragraph(t).paragraph_format.left_indent=Cm(0.5)

    # ====== 二、实验环境 ======
    doc.add_heading('二、实验环境',1)
    doc.add_paragraph('CPU: Intel Core i9-14900HX (8P+16E cores, 32 threads, base 2.2GHz)')
    doc.add_paragraph('内存: 64GB DDR5-5600')
    doc.add_paragraph('操作系统: Windows 11 + WSL2 (Ubuntu 24.04, Linux 5.15)')
    doc.add_paragraph('编译器: GCC 13.3.0, 编译选项: -O2')
    doc.add_paragraph('Cache参数 (getconf验证): L1d=48KB, L1i=32KB, L2=2MB/核, L3=36MB共享, Cache Line=64B')

    # ====== 三、实验内容和步骤 ======
    doc.add_heading('三、实验内容和步骤',1)

    # --- 3.1 矩阵乘法 ---
    doc.add_heading('3.1 分析Cache访存模式对系统性能的影响',2)
    doc.add_paragraph('通过对比两种不同循环顺序的矩阵乘法，分析访存模式（特别是空间局部性）对Cache效率的影响。')

    doc.add_heading('代码A：一般矩阵乘法（i-j-k循环）',3)
    add_code_block(doc, r'''// 一般矩阵乘法：i-j-k循环
// b[k*size+j] 每次访问步长 = size，空间局部性差
for(i = 0; i < size; i++) {
    for(j = 0; j < size; j++) {
        c[i * size + j] = 0;
        for(k = 0; k < size; k++)
            c[i * size + j] += a[i * size + k] * b[k * size + j];
    }
}

// 矩阵乘法对应的等价格式（与上述代码数学等价）：
for(j = 0; j < size; j++) {
    for(k = 0; k < size; k++) {
        temp = b[k * size + j];
        for(i = 0; i < size; i++)
            c[i * size + j] += a[i * size + k] * temp;
    }
}''')

    doc.add_paragraph('分析：上述代码中，b[k*size+j] 以步长size跨行访问（列优先访问），当size较大时（如≥1000），每次访问b[]都可能跨越多个Cache Line，导致Cache Miss率极高。')

    doc.add_heading('优化算法（i-k-j循环）',3)
    add_code_block(doc, r'''// 优化矩阵乘法：i-k-j循环
// 所有数组访问步长均为1，空间局部性极佳
for(i = 0; i < size; i++)
    for(j = 0; j < size; j++)
        c[i * size + j] = 0;    // 清零

for(i = 0; i < size; i++) {
    for(k = 0; k < size; k++) {
        temp = a[i * size + k];        // a[][] 步长=1
        for(j = 0; j < size; j++)
            c[i * size + j] += temp * b[k * size + j];
            // c[][] 步长=1, b[][] 步长=1  ← 关键优化！
    }
}''')

    doc.add_paragraph('优化要点：将b[]的访问从列优先（stride=size）改为行优先（stride=1），使所有三个数组的访问步长均为1，最大化空间局部性。在数据规模较大时，这一优化可带来10倍以上的性能提升。')

    doc.add_page_break()

    # --- 3.2 Memory Mountain ---
    doc.add_heading('3.2 测量X86机器上的Cache层次结构和容量',2)

    doc.add_heading('测量原理',3)
    doc.add_paragraph('核心思想：通过改变访问数组大小（控制工作集大小）和访问步长（控制空间局部性），测量读吞吐量的变化。当工作集大小超过某级Cache容量时，读吞吐量会出现明显下降（"悬崖"），从而定位各级Cache的容量边界。')

    doc.add_paragraph('关键公式：')
    doc.add_paragraph('    读吞吐量 (MB/s) = (size / stride) × CPU频率 / 时钟周期数')
    doc.add_paragraph('其中 size/stride 表示有效读取的字节数。')

    doc.add_heading('完整测量程序代码',3)
    # Load the actual code
    with open(os.path.join(BASE_DIR, 'wsl_mountain.c'), encoding='utf-8') as f:
        mountain_code = f.read()
    add_code_block(doc, mountain_code)

    doc.add_paragraph('编译运行：')
    doc.add_paragraph('    gcc -O2 wsl_mountain.c -o mountain -lm')
    doc.add_paragraph('    ./mountain > mountain_data.txt')
    doc.add_paragraph('    然后用 Python/matplotlib 进行 3D 可视化。')

    doc.add_page_break()

    # --- 3.3 TLB ---
    doc.add_heading('3.3 测量X86机器TLB的大小（选做）',2)
    doc.add_paragraph('TLB（Translation Lookaside Buffer）是MMU中的地址翻译缓存。通过固定小步长(stride=1)、逐渐增大数组大小（增大页表压力），观察吞吐量下降的位置来推断TLB覆盖范围。')
    doc.add_paragraph('当数组大小超过TLB能覆盖的地址空间时，每次访存都可能触发TLB Miss，需要逐级查页表（Page Table Walk），导致吞吐量进一步下降。')

    # ====== 四、实验结果及分析 ======
    doc.add_heading('四、实验结果及分析',1)

    # --- 4.1 矩阵乘法 ---
    doc.add_heading('4.1 分析Cache访存模式对系统性能的影响',2)

    doc.add_heading('表1：普通矩阵乘法与优化矩阵乘法的性能对比（WSL2实测）',3)

    # 表格
    tbl = doc.add_table(rows=9, cols=5); tbl.style='Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['矩阵大小\n(N×N)', '一般算法 (s)', '优化算法 (s)', '加速比', '备注']
    for j,h in enumerate(headers):
        c = tbl.cell(0,j); c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(9)
        set_cell_shading(c, '2C3E50')
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    rows_data = [
        ('100',  '0.000255',  '0.000278',  '0.92x',  '数据全部在L1中，优化开销>收益'),
        ('500',  '0.046886',  '0.033446',  '1.40x',  '数据在L2内，优化开始生效'),
        ('1000', '0.407510',  '0.252517',  '1.61x',  '数据部分溢出L2'),
        ('1500', '1.432449',  '0.866728',  '1.65x',  'L2→L3过渡区'),
        ('2000', '11.117311', '2.111857',  '5.26x',  '★ 一般算法超过L3(36MB)'),
        ('2500', '40.267149', '4.502034',  '8.94x',  '★ 加速比急剧增大'),
        ('3000', '97.751168', '8.437031',  '11.59x', '★ 最大加速比11.6倍'),
    ]
    for i,row in enumerate(rows_data):
        for j,val in enumerate(row):
            c=tbl.cell(i+1,j); c.text=''
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val).font.size=Pt(9)

    # 加速比公式
    tbl.cell(8,0).merge(tbl.cell(8,4))
    c=tbl.cell(8,0); c.text=''
    p=c.paragraphs[0]; r=p.add_run('加速比 = T_naive / T_optimized。加速比>1表示优化有效；加速比越大，Cache优化效果越显著。')
    r.font.size=Pt(9); r.italic=True

    doc.add_paragraph()
    add_img(doc, IMG['fig1_1'], Inches(5.5), '图1.1：一般算法与优化算法的执行时间对比（左：线性坐标，右：对数坐标）')
    add_img(doc, IMG['fig1_2'], Inches(5.3), '图1.2：加速比随矩阵大小的变化关系')

    doc.add_heading('原因分析',3)
    analyses = [
        '1. 小矩阵（N≤500）：加速比仅1.0x~1.4x。此时两个矩阵的数据量（500²×4×3≈3MB）仍在L3 Cache（36MB）范围内，且i9-14900HX的硬件预取器(IP Prefetcher)有效缓解了非连续访问的开销。优化效果不显著。',
        '2. 中等矩阵（N=1000~1500）：加速比1.6x。三个矩阵共需~27MB（1500²×4×3=27MB），接近L3边界但尚未超出。一般算法的b[]列访问开始遇到显著的Cache Line颠簸（thrashing），但优化算法保持良好空间局部性。',
        '3. 大矩阵（N=2000~3000）：加速比从5.3x跃升至11.6x！这是本实验最关键的发现。此时三个矩阵共需48~108MB，远超L3 Cache（36MB）。一般算法（i-j-k）的b[k*size+j]列访问步长=size，每次访问几乎都发生Cache Miss（→主存延迟~100ns），而优化算法（i-k-j）保持步长=1的连续访问，受益于Cache Line预取和突发传输，性能优势急剧扩大。',
        '4. 加速比在N=3000时达到11.59x且仍在增长趋势中，预测更大矩阵下加速比可达15x以上，直至内存带宽成为二者共同的瓶颈。',
        '5. 本实验数据比教材示例（加速比最大~2.5x）更加显著，原因在于：(a) 本机L3 Cache（36MB）远大于教材示例机器；(b) 现代CPU的预取器对不同访存模式的放大效应更为明显；(c) 使用了-O2编译优化。',
    ]
    for t in analyses:
        doc.add_paragraph(t)

    doc.add_page_break()

    # --- 4.2 Cache层次结构 ---
    doc.add_heading('4.2 测量分析Cache的层次结构、容量及L1 Cache行数',2)

    doc.add_heading('4.2.1 Memory Mountain可视化结果',3)
    doc.add_paragraph('运行测量程序后，获得38个不同数组大小（16KB~63MB）×16个步长（1~16）的读吞吐量数据，可视化如下：')

    add_img(doc, IMG['fig2_1'], Inches(5.0), '图2.1：Memory Mountain三维曲面图（真实WSL2测量数据）')
    add_img(doc, IMG['fig2_2'], Inches(5.0), '图2.2：Memory Mountain等高线图（标注了L1d/L2/L3 Cache边界）')

    doc.add_heading('4.2.2 Cache层次结构分析',3)
    doc.add_paragraph('从Memory Mountain图可以清晰观察到四级存储器层次：')

    # Cache表
    ct = doc.add_table(rows=5, cols=5); ct.style='Table Grid'
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    ch = ['存储层级', '容量', '步长=1吞吐量', '步长=8吞吐量', '测量方法']
    for j,h in enumerate(ch):
        c=ct.cell(0,j); c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(9)
        set_cell_shading(c,'2C3E50')
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

    cd = [
        ('L1 Data Cache','48 KB / 核','~76 GB/s','~66 GB/s','size≤48KB时吞吐最高'),
        ('L2 Cache','2 MB / 核','~62 GB/s','~34 GB/s','48KB<size≤2MB'),
        ('L3 Cache','36 MB 共享','~48 GB/s','~14 GB/s','2MB<size≤36MB'),
        ('Main Memory','DDR5 64GB','~21 GB/s','~4.5 GB/s','size>36MB'),
    ]
    for i,row in enumerate(cd):
        for j,val in enumerate(row):
            c=ct.cell(i+1,j); c.text=''
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val).font.size=Pt(9)
    doc.add_paragraph()

    add_img(doc, IMG['fig2_3'], Inches(5.3), '图2.3：Stride=1和Stride=8时吞吐量随Array Size的变化（Cache边界清晰可见）')

    doc.add_paragraph('关键观察：')
    doc.add_paragraph('  • 在 ~48KB 处出现第一个吞吐量拐点 → L1d Cache边界（48KB）确认')
    doc.add_paragraph('  • 在 ~2MB 处出现第二个吞吐量拐点 → L2 Cache边界（2MB）确认')
    doc.add_paragraph('  • 在 ~36MB 处出现第三个吞吐量拐点 → L3 Cache边界（36MB）确认')
    doc.add_paragraph('  • 大于36MB后吞吐量降至~21 GB/s并趋于稳定 → 主存带宽')

    doc.add_paragraph('结论：本机具有三级Cache（L1/L2/L3），容量分别为48KB、2MB、36MB。')

    doc.add_heading('4.2.3 L1 Cache行大小分析',3)
    add_img(doc, IMG['fig2_4'], Inches(5.3), '图2.4：不同固定Array Size下吞吐量随Stride的变化（确定Cache Block大小）')

    doc.add_paragraph('从图2.4观察到：')

    block_analysis = [
        '步长从1增加到7时，吞吐量缓慢下降（由于每次访问跳过更多有效数据）；',
        '步长=8时出现明显的吞吐量断崖——跨过了64字节Cache Line边界，每次访问都需要加载新的Cache Line；',
        '步长>8后吞吐量趋于稳定——每个Cache Line中只使用了一个long（8字节），其余56字节被浪费。',
        '',
        '因此：Cache Block（Line） = 8 elements × 8 bytes/element = 64 Bytes',
        '',
        'L1 Data Cache行数计算：',
        '  行数 = L1d容量 / Block大小 = 48 KB / 64 B = 49,152 B / 64 B = 768 行',
        '  组数 = 768 / 12-way = 64 组（12路组相联）',
    ]
    for t in block_analysis:
        doc.add_paragraph(t)

    doc.add_heading('4.2.4 验证实验结果（getconf）',3)
    doc.add_paragraph('在WSL2 Linux环境中运行 getconf 命令验证：')

    add_code_block(doc, r'''$ getconf -a | grep CACHE
LEVEL1_ICACHE_SIZE                 32768      (32 KB)
LEVEL1_ICACHE_LINESIZE             64
LEVEL1_DCACHE_SIZE                 49152      (48 KB)  ← 与Memory Mountain一致
LEVEL1_DCACHE_ASSOC                12
LEVEL1_DCACHE_LINESIZE             64                   ← Cache Line 64B确认
LEVEL2_CACHE_SIZE                  2097152    (2 MB)    ← 与Memory Mountain一致
LEVEL2_CACHE_ASSOC                 16
LEVEL2_CACHE_LINESIZE              64
LEVEL3_CACHE_SIZE                  37748736   (36 MB)   ← 与Memory Mountain一致
LEVEL3_CACHE_ASSOC                 12
LEVEL3_CACHE_LINESIZE              64''')

    add_img(doc, IMG['fig2_5'], Inches(5.3), '图2.5：Cache验证 - getconf命令输出与实验测量结果对比')

    doc.add_paragraph('验证结论：getconf命令输出的Cache参数与Memory Mountain实验测量结果完全吻合，验证了实验方法和分析的正确性。')

    doc.add_page_break()

    # --- 4.3 TLB ---
    doc.add_heading('4.3 TLB大小测量分析（选做）',2)
    doc.add_paragraph('通过固定stride=1、逐步增大数组大小的方式测量TLB覆盖范围。当数组大小超过TLB覆盖的地址空间时，TLB Miss率上升，需要访问内存中的页表（Page Table Walk），吞吐量会出现额外下降。')
    doc.add_paragraph('在本机（i9-14900HX）上，Intel手册给出的TLB参数为：')
    doc.add_paragraph('  • L1 iTLB: 64 entries (4KB pages) → covers 256KB')
    doc.add_paragraph('  • L1 dTLB: 96 entries (4KB pages) → covers 384KB')
    doc.add_paragraph('  • L2 STLB: 2048 entries (4KB pages) → covers 8MB')
    doc.add_paragraph('在Memory Mountain图中，stride=1时在8MB附近未观察到额外下降（L3的36MB边界掩盖了STLB边界），这是因为L3 Cache Miss的开销远大于STLB Miss，在图中无法区分。要精确测量TLB需要更精细的实验设计（如使用随机访问模式消除Cache预取的影响）。')

    # ====== 五、实验结论 ======
    doc.add_heading('五、实验结论',1)

    conclusions = [
        '1. 空间局部性与程序性能：',
        '   通过矩阵乘法优化实验，亲眼见证了访存模式对程序性能的巨大影响。相同时间复杂度的两个算法，仅因循环顺序不同（i-j-k vs i-k-j），在N=3000时性能差距达到11.6倍。这深刻说明了即使算法"正确"，如果访存模式与Cache结构不匹配，实际性能可能严重劣化。优化后所有数组访问步长均为1，完美利用Cache Line预取机制，将原本几乎100%的Cache Miss率降低到接近理论最优水平。',

        '2. Cache层次结构的测量：',
        '   通过Memory Mountain实验，成功使用软件方法"透视"了底层硬件——从读吞吐量的三个明显下降台阶（48KB、2MB、36MB），精确确定了本机三级Cache的容量。更令人印象深刻的是，这些测量值与getconf命令输出的硬件参数完全一致，验证了CSAPP教材所述方法的正确性。从3D可视化图表中，四个明显的"山脊"直观展示了L1→L2→L3→主存的速度递减关系。',

        '3. Cache Block大小的推断：',
        '   通过分析吞吐量随步长的变化，发现stride=8是性能陡降的临界点（8×8=64字节），精确确定了Cache Block=64字节。进一步计算出L1d Cache有768行、分为64组（12路组相联）。这种通过软件推演硬件参数的过程，加深了对组相联Cache映射机制的理解。',

        '4. 实验心得：',
        '   本次实验让我深刻体会到"Memory Wall"问题的现实意义。尽管CPU频率在过去几十年增长了数千倍，但内存延迟的改善仅为数十倍。Cache作为弥合这一差距的关键技术，其重要性不言而喻。在实际编程中，我将遵循以下准则：',
        '     • 优先使用连续内存访问模式（stride=1）',
        '     • 循环嵌套顺序应使最内层循环访问连续内存',
        '     • 关注关键数据的工作集大小，合理分块（Blocking）以适配Cache容量',
        '     • 在性能关键代码中，可通过微基准测试验证Cache行为',
        '   这些原则对高性能计算、数据库系统、游戏引擎等领域的代码优化都具有普适意义。',
    ]

    for t in conclusions:
        if t.endswith('：'):
            p = doc.add_paragraph(t)
            p.runs[0].bold = True
        else:
            doc.add_paragraph(t)

    doc.add_paragraph()
    doc.add_paragraph()

    # 评阅
    doc.add_heading('指导教师批阅意见：',2)
    doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('指导教师签字：').font.size=Pt(12)
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('2026年    月    日').font.size=Pt(12)

    # 保存
    doc.save(OUTPUT)
    print(f'\nReport saved to: {OUTPUT}')

if __name__ == '__main__':
    build()
